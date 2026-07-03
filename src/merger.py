import os
import shutil
import tempfile
import string
import re
import subprocess
import platform
import time
import threading
import queue
import concurrent.futures
from src.config import load_config
from src.filters import GitIgnoreFilter, _get_ignore_config

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from docx2pdf import convert as convert_docx
    DOCX2PDF_SUPPORT = True
except ImportError:
    DOCX2PDF_SUPPORT = False


def _get_libreoffice_path():
    system = platform.system()
    if system == "Windows":
        p = r"C:\Program Files\LibreOffice\program\soffice.exe"
        if os.path.exists(p):
            return p
    elif system == "Darwin":
        p = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if os.path.exists(p):
            return p
    which_soffice = shutil.which("soffice")
    if which_soffice:
        return which_soffice
    return "soffice"


def _run_with_cancel(cmd, cancel_event, timeout=300):
    try:
        proc = subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        while proc.poll() is None:
            if cancel_event.is_set():
                proc.terminate()
                try:
                    proc.wait(timeout=5)
                except subprocess.TimeoutExpired:
                    proc.kill()
                return False
            time.sleep(0.1)
        return proc.returncode == 0
    except Exception:
        return False


def _extract_legacy_doc_binary(file_path):
    """Brute-force extracts printable text from a legacy .doc binary file."""
    try:
        with open(file_path, "rb") as f:
            data = f.read()

        cleaned_data = data.replace(b'\x00', b'')
        raw_text = cleaned_data.decode('utf-8', errors='ignore')

        printable = set(string.printable)
        filtered_text = ''.join(filter(lambda x: x in printable, raw_text))

        filtered_text = re.sub(r'\n\s*\n', '\n\n', filtered_text)
        filtered_text = re.sub(r' {2,}', ' ', filtered_text)

        return filtered_text.strip()
    except Exception as e:
        return f"[Failed to extract legacy .doc text: {e}]"


def _extract_text(file_path, kind, log_callback=None, display_name=None):
    if kind == 'docx' and DOCX_SUPPORT:
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif kind == 'doc':
        return _extract_legacy_doc_binary(file_path)
    else:
        with open(file_path, "r", encoding="utf-8", errors="replace") as infile:
            return infile.read()


def _copy_large_file(outfile, task, log_callback=None):
    outfile.write(f"----- {task.display_name} -----\n")
    input_buffer = 128 * 1024
    if task.kind == 'docx' and DOCX_SUPPORT:
        doc = docx.Document(task.path)
        for para in doc.paragraphs:
            outfile.write(para.text + "\n")
    elif task.kind == 'doc':
        outfile.write(_extract_legacy_doc_binary(task.path))
    else:
        with open(task.path, 'r', encoding='utf-8', errors='replace', buffering=input_buffer) as infile:
            shutil.copyfileobj(infile, outfile)
    outfile.write("\n")
    if log_callback:
        log_callback(f"Merged large file: {task.display_name}")


def _parallel_text_merge(tasks, out_path, max_workers, cancel_event, large_file_threshold, progress_cb, log_callback, tree_str=None):
    result_queue = queue.PriorityQueue()
    pending = {}
    next_write_idx = 0

    small_tasks = [t for t in tasks if t.size < large_file_threshold]

    def worker(task):
        if cancel_event.is_set():
            return
        try:
            content = _extract_text(task.path, task.kind, log_callback, task.display_name)
            result_queue.put((task.index, task.display_name, content, True))
        except Exception as e:
            err_msg = f"[Error reading file: {e}]"
            if log_callback:
                log_callback(f"Error reading {task.display_name}: {e}")
            result_queue.put((task.index, task.display_name, err_msg, False))

    output_buffer = 256 * 1024

    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(worker, t) for t in small_tasks]

        with open(out_path, 'w', encoding='utf-8', buffering=output_buffer) as outfile:
            if tree_str:
                outfile.write("Directory Structure:\n")
                outfile.write(tree_str)
                outfile.write("\n\n--- File Contents ---\n\n")
            while next_write_idx < len(tasks):
                if cancel_event.is_set():
                    for f in futures:
                        f.cancel()
                    break

                current_task = tasks[next_write_idx]
                if current_task.size >= large_file_threshold:
                    if log_callback:
                        log_callback(f"Streaming large file: {current_task.display_name}")
                    try:
                        _copy_large_file(outfile, current_task, log_callback)
                    except Exception as e:
                        outfile.write(f"----- {current_task.display_name} -----\n")
                        outfile.write(f"[Error reading file: {e}]\n")
                        if log_callback:
                            log_callback(f"Error reading {current_task.display_name}: {e}")
                    next_write_idx += 1
                    if progress_cb:
                        progress_cb()
                    continue

                try:
                    idx, name, content, success = result_queue.get(timeout=0.05)
                    pending[idx] = (name, content, success)
                except queue.Empty:
                    pass

                while next_write_idx in pending:
                    name, content, success = pending.pop(next_write_idx)
                    outfile.write(f"----- {name} -----\n")
                    outfile.write(content)
                    outfile.write("\n")
                    if log_callback and success:
                        log_callback(f"Merged: {name}")
                    next_write_idx += 1
                    if progress_cb:
                        progress_cb()


def _sequential_text_merge(tasks, out_path, cancel_event, large_file_threshold, progress_cb, log_callback, tree_str=None):
    output_buffer = 256 * 1024
    with open(out_path, 'w', encoding='utf-8', buffering=output_buffer) as outfile:
        if tree_str:
            outfile.write("Directory Structure:\n")
            outfile.write(tree_str)
            outfile.write("\n\n--- File Contents ---\n\n")
        for task in tasks:
            if cancel_event.is_set():
                break

            if task.size >= large_file_threshold:
                if log_callback:
                    log_callback(f"Streaming large file: {task.display_name}")
                try:
                    _copy_large_file(outfile, task, log_callback)
                except Exception as e:
                    outfile.write(f"----- {task.display_name} -----\n")
                    outfile.write(f"[Error reading file: {e}]\n")
                    if log_callback:
                        log_callback(f"Error reading {task.display_name}: {e}")
            else:
                try:
                    content = _extract_text(task.path, task.kind, log_callback, task.display_name)
                    outfile.write(f"----- {task.display_name} -----\n")
                    outfile.write(content)
                    outfile.write("\n")
                    if log_callback:
                        log_callback(f"Merged: {task.display_name}")
                except Exception as e:
                    outfile.write(f"----- {task.display_name} -----\n")
                    outfile.write(f"[Error reading file: {e}]\n")
                    if log_callback:
                        log_callback(f"Error reading {task.display_name}: {e}")

            if progress_cb:
                progress_cb()


def _batch_libreoffice_convert(tasks, pdf_temp_dir, cancel_event, log_callback):
    if not tasks:
        return {}

    lo_path = _get_libreoffice_path()
    if lo_path == "soffice" and not shutil.which("soffice"):
        if log_callback:
            log_callback("LibreOffice not found in PATH. Skipping batch conversion.")
        return {}

    batch_dir = tempfile.mkdtemp()
    mapping = {}
    copied_paths = []

    try:
        for t in tasks:
            if cancel_event.is_set():
                return {}
            safe_basename = t.display_name.replace(os.sep, "_").replace("/", "_").replace("\\", "_")
            temp_copy_path = os.path.join(batch_dir, safe_basename)
            shutil.copy2(t.path, temp_copy_path)
            copied_paths.append(temp_copy_path)

            base_no_ext, _ = os.path.splitext(safe_basename)
            expected_pdf_name = base_no_ext + ".pdf"
            mapping[expected_pdf_name] = (t, temp_copy_path)

        cmd = [lo_path, "--headless", "--convert-to", "pdf", "--outdir", batch_dir] + copied_paths

        if log_callback:
            log_callback(f"Running batch LibreOffice conversion for {len(tasks)} files...")

        success = _run_with_cancel(cmd, cancel_event)
        if not success:
            if log_callback:
                log_callback("Batch LibreOffice conversion failed or was cancelled.")
            return {}

        results = {}
        for pdf_name, (task, _) in mapping.items():
            pdf_path_in_batch = os.path.join(batch_dir, pdf_name)
            if os.path.exists(pdf_path_in_batch):
                final_pdf_name = task.display_name.replace(os.sep, "_").replace("/", "_").replace("\\", "_") + ".pdf"
                final_pdf_path = os.path.join(pdf_temp_dir, final_pdf_name)

                shutil.move(pdf_path_in_batch, final_pdf_path)
                results[task.index] = final_pdf_path
                if log_callback:
                    log_callback(f"LibreOffice conversion successful: {task.display_name}")
            else:
                if log_callback:
                    log_callback(f"LibreOffice failed to convert: {task.display_name}")

        return results

    finally:
        shutil.rmtree(batch_dir, ignore_errors=True)


def _single_doc_convert(task, pdf_temp_dir, cancel_event, log_callback):
    is_docx = task.path.lower().endswith('.docx')
    final_pdf_name = task.display_name.replace(os.sep, "_").replace("/", "_").replace("\\", "_") + ".pdf"
    final_pdf_path = os.path.join(pdf_temp_dir, final_pdf_name)

    if is_docx and DOCX2PDF_SUPPORT:
        try:
            if log_callback:
                log_callback(f"Trying MS Word conversion for: {task.display_name}")
            convert_docx(task.path, final_pdf_path)
            if os.path.exists(final_pdf_path):
                return final_pdf_path
        except Exception as e:
            if log_callback:
                log_callback(f"MS Word conversion failed for {task.display_name}: {e}")

    lo_path = _get_libreoffice_path()
    has_lo = lo_path != "soffice" or shutil.which("soffice")
    if has_lo:
        batch_dir = tempfile.mkdtemp()
        try:
            safe_basename = task.display_name.replace(os.sep, "_").replace("/", "_").replace("\\", "_")
            temp_copy_path = os.path.join(batch_dir, safe_basename)
            shutil.copy2(task.path, temp_copy_path)

            cmd = [lo_path, "--headless", "--convert-to", "pdf", "--outdir", batch_dir, temp_copy_path]
            if log_callback:
                log_callback(f"Trying LibreOffice conversion for: {task.display_name}")

            if _run_with_cancel(cmd, cancel_event):
                base_no_ext, _ = os.path.splitext(safe_basename)
                expected_pdf_name = base_no_ext + ".pdf"
                pdf_path_in_batch = os.path.join(batch_dir, expected_pdf_name)
                if os.path.exists(pdf_path_in_batch):
                    shutil.move(pdf_path_in_batch, final_pdf_path)
                    return final_pdf_path
        except Exception as e:
            if log_callback:
                log_callback(f"LibreOffice single conversion failed for {task.display_name}: {e}")
        finally:
            shutil.rmtree(batch_dir, ignore_errors=True)

    return None


def _process_pdf_merge(
    tasks,
    pdf_temp_dir,
    pdf_list,
    styled_pdf,
    keep_txt_sources,
    txt_temp_dir,
    cancel_event,
    log_callback,
    item_callback,
    max_workers,
    perf_config
):
    pdf_paths_map = {}

    styled_tasks = []
    if styled_pdf:
        styled_tasks = [t for t in tasks if t.kind in ('docx', 'doc')]

    batch_enabled = perf_config.get("batch_libreoffice", True)
    failed_styled_tasks = []

    if styled_tasks:
        if batch_enabled:
            batch_results = _batch_libreoffice_convert(styled_tasks, pdf_temp_dir, cancel_event, log_callback)
            for t in styled_tasks:
                if t.index in batch_results:
                    pdf_paths_map[t.index] = batch_results[t.index]
                    if item_callback:
                        item_callback()
                else:
                    failed_styled_tasks.append(t)
        else:
            failed_styled_tasks = styled_tasks

        still_failed_styled = []
        for t in failed_styled_tasks:
            if cancel_event.is_set():
                return
            res_path = _single_doc_convert(t, pdf_temp_dir, cancel_event, log_callback)
            if res_path:
                pdf_paths_map[t.index] = res_path
                if item_callback:
                    item_callback()
            else:
                still_failed_styled.append(t)
        failed_styled_tasks = still_failed_styled

    plain_tasks = []
    if not styled_pdf:
        plain_tasks = tasks
    else:
        plain_tasks = [t for t in tasks if t.kind not in ('docx', 'doc')] + failed_styled_tasks

    plain_tasks.sort(key=lambda t: t.index)

    conversions = []
    temp_txt_files = []

    for t in plain_tasks:
        if cancel_event.is_set():
            break

        final_pdf_name = t.display_name.replace(os.sep, "_").replace("/", "_").replace("\\", "_") + ".pdf"
        final_pdf_path = os.path.join(pdf_temp_dir, final_pdf_name)

        is_temp_txt = False
        src_txt_path = t.path

        if t.kind in ('docx', 'doc') or keep_txt_sources:
            try:
                content = _extract_text(t.path, t.kind, log_callback, t.display_name)
            except Exception as e:
                content = f"[Error reading file: {e}]"
                if log_callback:
                    log_callback(f"Error reading {t.display_name}: {e}")

            if keep_txt_sources:
                safe_txt_name = t.display_name.replace(os.sep, "_").replace("/", "_").replace("\\", "_")
                if not safe_txt_name.lower().endswith(".txt"):
                    safe_txt_name = os.path.splitext(safe_txt_name)[0] + ".txt"
                persistent_txt_path = os.path.join(txt_temp_dir, safe_txt_name)
                try:
                    with open(persistent_txt_path, "w", encoding="utf-8") as tf:
                        tf.write(content)
                    if t.kind not in ('docx', 'doc'):
                        src_txt_path = persistent_txt_path
                except Exception as e:
                    if log_callback:
                        log_callback(f"Failed to save source txt for {t.display_name}: {e}")

            if t.kind in ('docx', 'doc'):
                temp_tf = tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8")
                temp_tf.write(content)
                temp_tf.close()
                src_txt_path = temp_tf.name
                is_temp_txt = True
                temp_txt_files.append(temp_tf.name)

        conversions.append((t.index, src_txt_path, final_pdf_path, t.display_name, is_temp_txt))

    if conversions and not cancel_event.is_set():
        parallel_pdf_fallback = perf_config.get("parallel_pdf_fallback", True)

        many_files = len(conversions) >= 20
        large_files = sum(tasks[idx].size for idx, _, _, _, _ in conversions) > 1 * 1024 * 1024
        use_processes = parallel_pdf_fallback and many_files and large_files

        worker_args = [(c[1], c[2], c[3], False) for c in conversions]

        if log_callback:
            mode_name = "ProcessPool" if use_processes else "ThreadPool"
            log_callback(f"Converting {len(conversions)} files to PDF in parallel using {mode_name}...")

        executor_cls = concurrent.futures.ProcessPoolExecutor if use_processes else concurrent.futures.ThreadPoolExecutor

        try:
            from src.pdf_utils import _convert_text_to_pdf_worker

            with executor_cls(max_workers=max_workers) as executor:
                future_to_index = {
                    executor.submit(_convert_text_to_pdf_worker, arg): conversions[i][0]
                    for i, arg in enumerate(worker_args)
                }

                for future in concurrent.futures.as_completed(future_to_index):
                    if cancel_event.is_set():
                        executor.shutdown(wait=False, cancel_futures=True)
                        break
                    idx = future_to_index[future]
                    try:
                        res_pdf_path = future.result()
                        pdf_paths_map[idx] = res_pdf_path
                        if item_callback:
                            item_callback()
                    except Exception as e:
                        task_disp = tasks[idx].display_name
                        if log_callback:
                            log_callback(f"Failed to convert {task_disp} to PDF: {e}")
        except Exception as e:
            if log_callback:
                log_callback(f"Parallel PDF conversion error: {e}")

    for temp_file in temp_txt_files:
        try:
            os.remove(temp_file)
        except Exception:
            pass

    for i in range(len(tasks)):
        if i in pdf_paths_map:
            pdf_list.append(pdf_paths_map[i])


def _cleanup_pdf_temp(pdf_temp_dir, pdf_list, keep_pdf_sources):
    if not keep_pdf_sources and pdf_temp_dir and os.path.exists(pdf_temp_dir):
        shutil.rmtree(pdf_temp_dir, ignore_errors=True)


def merge_files(
    directory,
    config=None,
    extension=None,
    recursive=False,
    output_file=None,
    ignore_dirs=None,
    ignore_exts=None,
    cancel_event=None,
    dry_run=False,
    log_callback=None,
    item_callback=None,
    use_gitignore=True,
    pdf_mode=False,
    keep_pdf_sources=False,
    keep_txt_sources=False,
    styled_pdf=False,
    tasks=None,
    is_git=False,
    git_branch=None,
    git_tag=None,
    git_commit=None,
    git_token=None,
    include_tree=None,
    tasks_collected_callback=None,
    include_list=None,
    tree_ignore_level=None
):
    if config is None:
        config = load_config()

    if cancel_event is None:
        cancel_event = threading.Event()

    perf = config.get("performance", {})
    max_workers = perf.get("max_workers", 0)
    if max_workers <= 0:
        max_workers = min(32, os.cpu_count() + 4)
    large_file_threshold_mb = perf.get("large_file_threshold_mb", 5)
    large_file_threshold = large_file_threshold_mb * 1024 * 1024
    min_tasks_for_parallel = perf.get("min_tasks_for_parallel", 8)
    pdf_batch_threshold = perf.get("pdf_batch_threshold", 200)

    # Git Ingestion logic
    temp_dir_to_clean = None
    git_url = None
    if is_git or (directory and (directory.startswith("http://") or directory.startswith("https://"))):
        is_git = True
        git_url = directory
        temp_dir_to_clean = tempfile.mkdtemp()
        try:
            if log_callback:
                log_callback(f"Cloning Git repository: {git_url}...")
            from src.git_utils import clone_repo
            clone_repo(
                url=git_url,
                target_dir=temp_dir_to_clean,
                branch=git_branch,
                tag=git_tag,
                commit=git_commit,
                token=git_token
            )
            directory = temp_dir_to_clean
        except Exception as e:
            shutil.rmtree(temp_dir_to_clean, ignore_errors=True)
            if log_callback:
                log_callback(f"Git clone failed: {e}")
            raise e

    try:
        raw_out_path = output_file or config.get("output_file", "Mono.txt")
        out_dir = config.get("output_dir", "out")
        out_path = os.path.join(out_dir, os.path.basename(raw_out_path))

        ignore_set, ignored_ext_tuple, ignored_files = _get_ignore_config(config, ignore_dirs, ignore_exts)
        git_filter = GitIgnoreFilter(directory) if use_gitignore else None

        if tasks is None:
            from src.collector import collect_files
            tasks = collect_files(
                directory=directory,
                extension=extension,
                recursive=recursive,
                ignore_set=ignore_set,
                ignored_ext_tuple=ignored_ext_tuple,
                ignored_files=ignored_files,
                git_filter=git_filter,
                include_list=include_list
            )
            if tasks_collected_callback:
                tasks_collected_callback(tasks)

        if not tasks:
            if log_callback:
                log_callback("No files found to process.")
            return None

        # Resolve include_tree parameter
        if include_tree is None:
            include_tree = config.get("include_tree", True)

        if tree_ignore_level is None:
            tree_ignore_level = config.get("tree_ignore_level", "none")

        tree_str = None
        if include_tree:
            from src.tree_utils import generate_tree
            if tree_ignore_level == "all":
                tree_tasks = tasks
            else:
                from src.collector import collect_files
                if tree_ignore_level == "extension":
                    t_ignore_set = ignore_set | {".git"}
                    t_ignored_ext_tuple = ignored_ext_tuple
                    t_ignored_files = ignored_files
                    t_extension = extension
                elif tree_ignore_level == "settings":
                    t_ignore_set = ignore_set | {".git"}
                    t_ignored_ext_tuple = ignored_ext_tuple
                    t_ignored_files = ignored_files
                    t_extension = None
                else:  # "none" or default
                    t_ignore_set = {".git"}
                    t_ignored_ext_tuple = ()
                    t_ignored_files = set()
                    t_extension = None

                tree_tasks = collect_files(
                    directory=directory,
                    extension=t_extension,
                    recursive=recursive,
                    ignore_set=t_ignore_set,
                    ignored_ext_tuple=t_ignored_ext_tuple,
                    ignored_files=t_ignored_files,
                    git_filter=git_filter,
                    include_list=None
                )
            tree_str = generate_tree(tree_tasks, directory)

        if dry_run:
            token_count = 0
            try:
                accumulated_text = []
                if include_tree and tree_str:
                    accumulated_text.append("Directory Structure:\n" + tree_str + "\n\n--- File Contents ---\n\n")
                for task in tasks:
                    if cancel_event.is_set():
                        break
                    try:
                        text_content = _extract_text(task.path, task.kind, log_callback, task.display_name)
                        accumulated_text.append(f"----- {task.display_name} -----\n{text_content}\n")
                    except Exception:
                        pass
                from src.token_utils import estimate_tokens
                token_count = estimate_tokens("".join(accumulated_text))
            except Exception as e:
                if log_callback:
                    log_callback(f"Failed to calculate preview token count: {e}")

            for task in tasks:
                if cancel_event.is_set():
                    break
                if log_callback:
                    log_callback(f"Would merge: {task.display_name}")
                if item_callback:
                    item_callback()
            return {
                "file_count": len(tasks),
                "total_size_bytes": sum(t.size for t in tasks),
                "token_count": token_count,
                "output_path": out_path,
                "tree": tree_str
            }

        if pdf_mode:
            base, _ = os.path.splitext(out_path)
            out_path = base + ".pdf"

        if is_git:
            repo_name = git_url.split("/")[-1]
            if repo_name.endswith(".git"):
                repo_name = repo_name[:-4]
            source_dir_name = repo_name
        else:
            source_dir_name = os.path.basename(os.path.normpath(directory))

        if not source_dir_name:
            source_dir_name = "merged_sources"

        pdf_temp_dir = None
        pdf_list = []
        if pdf_mode:
            if keep_pdf_sources:
                pdf_temp_dir = os.path.join(out_dir, source_dir_name, "pdf")
            else:
                base_filename = os.path.splitext(os.path.basename(out_path))[0]
                pdf_temp_dir = os.path.join(out_dir, base_filename)
            os.makedirs(pdf_temp_dir, exist_ok=True)

        txt_temp_dir = None
        if keep_txt_sources:
            txt_temp_dir = os.path.join(out_dir, source_dir_name, "txt")
            os.makedirs(txt_temp_dir, exist_ok=True)

        os.makedirs(out_dir, exist_ok=True)
        tmp_path = out_path + ".tmp"

        if pdf_mode:
            _process_pdf_merge(
                tasks=tasks,
                pdf_temp_dir=pdf_temp_dir,
                pdf_list=pdf_list,
                styled_pdf=styled_pdf,
                keep_txt_sources=keep_txt_sources,
                txt_temp_dir=txt_temp_dir,
                cancel_event=cancel_event,
                log_callback=log_callback,
                item_callback=item_callback,
                max_workers=max_workers,
                perf_config=perf
            )

            if cancel_event.is_set():
                if log_callback:
                    log_callback("Operation cancelled.")
                _cleanup_pdf_temp(pdf_temp_dir, pdf_list, keep_pdf_sources)
                return None

            if not pdf_list:
                if log_callback:
                    log_callback("No PDF sources generated.")
                return None

            # Generate and prepend directory tree in PDF mode
            if include_tree and tree_str:
                tree_txt_fd, tree_txt_path = tempfile.mkstemp(suffix=".txt")
                try:
                    with os.fdopen(tree_txt_fd, "w", encoding="utf-8") as tf:
                        tf.write("Directory Structure:\n")
                        tf.write(tree_str)
                        tf.write("\n")
                    tree_pdf_path = os.path.join(pdf_temp_dir, "_directory_structure.pdf")
                    from src.pdf_utils import convert_to_pdf
                    convert_to_pdf(tree_txt_path, tree_pdf_path, "Directory Structure", styled=styled_pdf)
                    pdf_list.insert(0, tree_pdf_path)
                finally:
                    try:
                        os.remove(tree_txt_path)
                    except Exception:
                        pass

            if log_callback:
                log_callback("Compiling final PDF structure...")

            from src.pdf_utils import _merge_pdf_files
            _merge_pdf_files(pdf_list, tmp_path, pdf_batch_threshold, log_callback)

            if cancel_event.is_set():
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
                _cleanup_pdf_temp(pdf_temp_dir, pdf_list, keep_pdf_sources)
                return None

            os.replace(tmp_path, out_path)

            if not keep_pdf_sources:
                shutil.rmtree(pdf_temp_dir, ignore_errors=True)
                if log_callback:
                    log_callback("Source files cleaned up completely.")
            else:
                if log_callback:
                    log_callback(f"Source files preserved in: {pdf_temp_dir}")

            # Calculate token count on accumulated source text
            accumulated_text = []
            if include_tree and tree_str:
                accumulated_text.append("Directory Structure:\n" + tree_str + "\n\n--- File Contents ---\n\n")
            for task in tasks:
                if cancel_event.is_set():
                    break
                try:
                    text_content = _extract_text(task.path, task.kind, log_callback, task.display_name)
                    accumulated_text.append(f"----- {task.display_name} -----\n{text_content}\n")
                except Exception:
                    pass
            from src.token_utils import estimate_tokens
            token_count = estimate_tokens("".join(accumulated_text))
            final_size = os.path.getsize(out_path)

        else:
            small_tasks = [t for t in tasks if t.size < large_file_threshold]
            if len(small_tasks) >= min_tasks_for_parallel:
                _parallel_text_merge(
                    tasks=tasks,
                    out_path=tmp_path,
                    max_workers=max_workers,
                    cancel_event=cancel_event,
                    large_file_threshold=large_file_threshold,
                    progress_cb=item_callback,
                    log_callback=log_callback,
                    tree_str=tree_str
                )
            else:
                _sequential_text_merge(
                    tasks=tasks,
                    out_path=tmp_path,
                    cancel_event=cancel_event,
                    large_file_threshold=large_file_threshold,
                    progress_cb=item_callback,
                    log_callback=log_callback,
                    tree_str=tree_str
                )

            if cancel_event.is_set():
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
                if log_callback:
                    log_callback("Operation cancelled.")
                return None

            os.replace(tmp_path, out_path)

            # Read merged output to compute final size and token count
            final_size = os.path.getsize(out_path)
            with open(out_path, "r", encoding="utf-8", errors="replace") as f:
                merged_content = f.read()
            from src.token_utils import estimate_tokens
            token_count = estimate_tokens(merged_content)

        res = {
            "file_count": len(tasks),
            "total_size_bytes": final_size,
            "token_count": token_count,
            "output_path": out_path,
            "tree": tree_str
        }
        return res

    except Exception as e:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass
        if pdf_mode:
            _cleanup_pdf_temp(pdf_temp_dir, pdf_list, keep_pdf_sources)
        raise e
    finally:
        if temp_dir_to_clean:
            shutil.rmtree(temp_dir_to_clean, ignore_errors=True)
